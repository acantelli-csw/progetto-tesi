from PIL import Image
from io import BytesIO
from docx import Document
import numpy as np
import tempfile
import os
import re
import fitz  
import win32com.client as win32

# Pulizia parti inutili del testo prima del chunking
def remove_placeholders(text: str) -> str:
    # Regex che elimina blocchi tipo [N.B. .... ] – -- -- testo da cancellare
    pattern = r"\[\s*N\.?B\.?.*?\]\s*[-–—\s]*"
    return re.sub(pattern, "", text, flags=re.IGNORECASE | re.DOTALL)


# Estrazione testo e trascrizione immagini dai documenti in base all'estensione
def extract_text_from_varbinary(file_data, extension, numero, reader):

    full_text = ""
    ext = "".join(extension.lower().split())

    try:

        # Caso 1: PDF (testo + OCR)
        if ext == ".pdf":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file_data)
                tmp_path = tmp.name

            # --- Testo nativo ---
            with fitz.open(tmp_path) as pdf:
                for page in pdf:
                    text = page.get_text()
                    if text:
                        full_text += text + "\n"

            # --- OCR su immagini ---
            with fitz.open(tmp_path) as doc:
                for page_index, page in enumerate(doc):
                    for img_index, img in enumerate(page.get_images(full=True)):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]

                        img = Image.open(BytesIO(image_bytes))
                        img_np = np.array(img)
                        result = reader.readtext(img_np)

                        ocr_text = " ".join([res[1] for res in result])

                        if ocr_text.strip():
                            full_text += f"\n[OCR p.{page_index+1} img.{img_index+1}]: {ocr_text}\n"

            os.remove(tmp_path)

        # Caso 2: DOCX (testo + OCR)
        elif ext == ".docx":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(file_data)
                tmp_path = tmp.name

            doc = Document(tmp_path)

            # --- Testo nativo ---
            for p in doc.paragraphs:
                full_text += p.text + "\n"

            # --- OCR su immagini embedded ---
            for rel in doc.part.rels.values():
                target_ref = getattr(rel, "target_ref", "")

                # Se la relazione punta a un'immagine
                if "image" in target_ref:
                    try:
                        target_part = getattr(rel, "target_part", None)
                        img_data = target_part.blob
                        img = Image.open(BytesIO(img_data))
                        img_np = np.array(img)
                        result = reader.readtext(img_np)

                        ocr_text = " ".join([res[1] for res in result])
                        if ocr_text.strip():
                            full_text += f"\n[OCR immagine embedded]: {ocr_text}\n"

                    except Exception as e:
                        print(f"Errore OCR: immagine esterna al DOCX. Elaboro il resto...\n")
                    
            os.remove(tmp_path)

        # Caso 3: DOC (converti in DOCX con Word e gestisci come DOCX)
        elif ext == ".doc":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
                tmp.write(file_data)
                tmp_doc_path = tmp.name

            # Percorso DOCX convertito
            tmp_docx_path = tmp_doc_path + ".docx"

            # --- Conversione tramite Word ---
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            doc_obj = word.Documents.Open(tmp_doc_path)

            # 16 = formato DOCX
            doc_obj.SaveAs(tmp_docx_path, FileFormat=16)
            doc_obj.Close()
            word.Quit()

            # Da qui lo elaboriamo come un DOCX
            doc = Document(tmp_docx_path)

            # --- Testo nativo ---
            for p in doc.paragraphs:
                full_text += p.text + "\n"

            # --- OCR immagini embedded ---
            for rel in doc.part.rels.values():
                target_ref = getattr(rel, "target_ref", "")

                if "image" in target_ref:
                    try:
                        target_part = getattr(rel, "target_part", None)
                        img_data = target_part.blob
                        img = Image.open(BytesIO(img_data))
                        img_np = np.array(img)
                        result = reader.readtext(img_np)

                        ocr_text = " ".join([res[1] for res in result])
                        if ocr_text.strip():
                            full_text += f"\n[OCR immagine embedded]: {ocr_text}\n"

                    except Exception as e:
                        print(f"Errore OCR: immagine esterna al DOCX. Elaboro il resto...\n")

            os.remove(tmp_doc_path)
            os.remove(tmp_docx_path)

        # Caso non gestito
        else:
            print(f"Estensione non gestita: {ext}")
            full_text = ""

        # Pulizia testo
        full_text = full_text.strip().replace("\x0c", "")
        full_text = remove_placeholders(full_text)
        return full_text

    except Exception as e:
        print(f"Errore elaborando file {numero}{ext}: {e}")
        return ""